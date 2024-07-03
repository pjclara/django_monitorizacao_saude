from docx import Document
from docx.shared import Pt

# Cria um documento
doc = Document()

# Adiciona o título
doc.add_heading('Tabela de Campos do Paciente', level=1)

# Adiciona uma tabela com 4 colunas
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'

# Define as colunas
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Campo'
hdr_cells[1].text = 'Tipo de Variável'
hdr_cells[2].text = 'Dependência'
hdr_cells[3].text = 'Valor Exemplo'

# Adiciona os dados na tabela
data = [
    ["_id", "ObjectId", "N/A", '"668343268835bd9428ae2f59"'],
    ["sns", "Inteiro", "N/A", "123456677"],
    ["nome", "String", "N/A", '"Paciente 1"'],
    ["dataNascimento", "Data", "N/A", '"1990-10-10T00:00:00Z"'],
    ["genero", "String", "N/A", '"Masculino"'],
    ["telefone", "String", "N/A", '"912123123"'],
    ["peso", "Double", "N/A", "123.0"],
    ["altura", "Double", "N/A", "123.0"],
    ["dispositivos", "Array", "N/A", "Veja abaixo para detalhes de dispositivos"],
    ["dispositivos.profissionais.id", "Inteiro", "dispositivos", "1"],
    ["dispositivos.profissionais.role", "String", "dispositivos", '"profissional"'],
    ["dispositivos.data_inicio", "Data", "dispositivos", '"2024-07-01T00:00:00Z"'],
    ["dispositivos.data_fim", "Data", "dispositivos", '"2024-07-03T00:00:00Z"'],
    ["dispositivos.ativo", "Boolean", "dispositivos", "false"],
    ["dispositivos.modelo", "String", "dispositivos", '"X1"'],
    ["dispositivos.fabricante", "String", "dispositivos", '"Health"'],
    ["dispositivos.numeroSerie", "Inteiro", "dispositivos", "12345"],
    ["dispositivos.descricao", "String", "dispositivos", '"Smartband"'],
    ["dispositivos.sinaisVitais", "Array", "dispositivos", "Veja abaixo para detalhes de sinais vitais"],
    ["dispositivos.sinaisVitais.maximo", "Double", "sinaisVitais", "100.0"],
    ["dispositivos.sinaisVitais.minimo", "Double", "sinaisVitais", "55.0"],
    ["dispositivos.sinaisVitais.tipo", "String", "sinaisVitais", '"Frequência Cardíaca"'],
    ["dispositivos.sinaisVitais.unidade", "String", "sinaisVitais", '"bpm"'],
    ["dispositivos.sinaisVitais.valores", "Array", "sinaisVitais", "Veja abaixo para detalhes de valores"],
    ["dispositivos.sinaisVitais.valores.valor", "Double", "valores", "106.0"],
    ["dispositivos.sinaisVitais.valores.data", "Data", "valores", '"2024-07-02T10:21:19Z"'],
    ["dispositivos.sinaisVitais.valores.alerta", "Boolean", "valores", "true"],
    ["dispositivos.sinaisVitais.valores.lida", "Boolean", "valores", "false"],
    ["dispositivos.sinaisVitais.valores.dataLida", "Data", "valores", "null"],
    ["dispositivos.sinaisVitais.ativo", "Boolean", "sinaisVitais", "false"],
    ["dispositivos.sinaisVitais.readingFrequency", "Inteiro", "sinaisVitais", "5"]
]

# Preenche a tabela com os dados
for row in data:
    row_cells = table.add_row().cells
    for i, item in enumerate(row):
        row_cells[i].text = str(item)

# Salva o documento
file_path = "/mnt/data/tabela_paciente.docx"
doc.save(file_path)

file_path
